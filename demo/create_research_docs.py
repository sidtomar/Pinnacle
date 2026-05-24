"""
Generate 4 internal medical research documents for the PinnacleIQ demo.
Showcases Agent Alpha's local file reading capability.

Output folder:  <project_root>/Research/
Files created:
  1. SGLT2_Empagliflozin_Clinical_Summary_Q4_2024.docx   (Word)
  2. Diabetes_Drug_Interaction_Database_2024.xlsx          (Excel)
  3. GLP1_Real_World_India_Evidence_INTERNAL.txt           (Text)
  4. PCOS_Management_Protocol_India_2024.docx              (Word)

Run:
  D:\\venv\\pinnacle\\Scripts\\python.exe demo/create_research_docs.py
"""
from pathlib import Path

OUT = Path(__file__).parent.parent / "Research"
OUT.mkdir(exist_ok=True)


# =============================================================================
# FILE 1  SGLT2 Clinical Summary  (.docx)
# =============================================================================
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.core_properties.author = "Dr. Prashant Agarwal, Medical Affairs, Mankind Pharma"

h = doc.add_heading("SGLT2 Inhibitors: Clinical Evidence Summary", level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Document Type: Internal Clinical Summary  |  Q4 2024")
doc.add_paragraph("Prepared by: Medical Affairs, Mankind Pharma  |  CONFIDENTIAL")
doc.add_paragraph("")

doc.add_heading("1. Overview", level=2)
doc.add_paragraph(
    "SGLT2 inhibitors (sodium-glucose co-transporter-2 inhibitors) reduce blood glucose by blocking "
    "glucose reabsorption in the proximal convoluted tubule. Key molecules: empagliflozin (Jardiance), "
    "dapagliflozin (Farxiga), canagliflozin (Invokana). Now indicated beyond T2DM — heart failure and "
    "CKD regardless of glycaemic status."
)

doc.add_heading("2. EMPA-REG OUTCOME — Key Efficacy Data", level=2)
for label, text in [
    ("Primary Endpoint:", "MACE reduced by 14% vs placebo (HR 0.86, 95% CI 0.74-0.99; p=0.04)."),
    ("CV Death:", "38% relative risk reduction (HR 0.62, 95% CI 0.49-0.77; p<0.001)."),
    ("HF Hospitalisation:", "35% relative risk reduction (HR 0.65, 95% CI 0.50-0.85; p=0.002)."),
    ("Renal Composite:", "39% RRR in renal progression endpoint."),
]:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(" " + text)

doc.add_heading("3. Indian Real-World Evidence (Mankind Internal, 2024)", level=2)
doc.add_paragraph(
    "Multi-centre registry, n=2,847 T2DM patients, 18 centres, 2022-2024:"
)
for item in [
    "Mean HbA1c reduction: 1.2% (empagliflozin 10 mg) vs 0.8% (standard care) at 6 months",
    "Body weight reduction: 2.8 kg vs 0.4 kg over 12 months",
    "eGFR preservation: 94% maintained eGFR >60 at 2 years vs 87% standard care",
    "Genital mycotic infections: 4.1% — manageable with standard hygiene counselling",
    "Highest benefit: patients with established CVD or CKD (eGFR 30-60)",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4. 2025 ESC / ADA-EASD Positioning", level=2)
doc.add_paragraph(
    "SGLT2 inhibitors hold Class I, Level A recommendation in 2025 ESC Heart Failure guidelines "
    "for all HF patients — HFrEF and HFpEF alike — regardless of diabetes status. ADA/EASD 2025 "
    "consensus recommends SGLT2i as first-line alongside metformin for any T2DM patient with HF, "
    "CKD, or established ASCVD."
)

doc.add_heading("5. Competitive Landscape & Mankind Positioning", level=2)
doc.add_paragraph(
    "Key competitors: Boehringer Ingelheim (Jardiance), AstraZeneca (Forxiga), J&J (Invokana). "
    "Mankind advantage: 30-40% MRP advantage, equivalent efficacy, Indian registry data to back claims. "
    "Target MRs to lead with the HF/CKD story — expands the prescriber base beyond diabetologists."
)

doc.save(str(OUT / "SGLT2_Empagliflozin_Clinical_Summary_Q4_2024.docx"))
print("[OK] SGLT2_Empagliflozin_Clinical_Summary_Q4_2024.docx")


# =============================================================================
# FILE 2  Drug Interaction Database  (.xlsx)
# =============================================================================
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Drug Interactions"

NAVY  = PatternFill(start_color="003087", end_color="003087", fill_type="solid")
RED   = PatternFill(start_color="F8D7DA", end_color="F8D7DA", fill_type="solid")
AMBER = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
GREEN = PatternFill(start_color="D4EDDA", end_color="D4EDDA", fill_type="solid")
WHITE_BOLD = Font(color="FFFFFF", bold=True, size=11)
CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
WRAP   = Alignment(wrap_text=True, vertical="top")

ws.merge_cells("A1:G1")
ws["A1"] = "MANKIND PHARMA  —  Diabetes Drug Interaction Database 2024"
ws["A1"].font = Font(bold=True, size=13, color="003087")
ws["A1"].alignment = Alignment(horizontal="center")
ws.row_dimensions[1].height = 28

ws.merge_cells("A2:G2")
ws["A2"] = "Confidential  |  Medical Affairs  |  Updated: October 2024  |  For MR Training Use Only"
ws["A2"].font = Font(italic=True, color="6C757D", size=9)
ws["A2"].alignment = Alignment(horizontal="center")

ws.append([])  # blank row 3

headers = ["Drug (Mankind)", "Interacting Drug / Class", "Mechanism", "Severity",
           "Clinical Effect", "Management", "Reference"]
ws.append(headers)
for col, _ in enumerate(headers, 1):
    c = ws.cell(row=4, column=col)
    c.font = WHITE_BOLD
    c.fill = NAVY
    c.alignment = CENTER

ROWS = [
    ["Empagliflozin (DYNAGLIP)", "Insulin / Sulphonylureas", "Additive glucose lowering", "MODERATE",
     "Increased hypoglycaemia risk", "Reduce insulin/SU dose 20-25% on initiation", "ADA 2024"],
    ["Empagliflozin (DYNAGLIP)", "Loop Diuretics (Furosemide)", "Additive diuresis/natriuresis", "MODERATE",
     "Dehydration, orthostatic hypotension", "Monitor BP + hydration; reduce diuretic if needed", "ESC HF 2023"],
    ["Empagliflozin (DYNAGLIP)", "ACE inhibitors / ARBs", "Complementary renoprotective", "MINOR",
     "Enhanced BP lowering (usually beneficial)", "Monitor eGFR + electrolytes at week 4", "CREDENCE 2023"],
    ["Metformin (MANKIND-M)", "Iodinated Contrast Media", "Lactic acidosis risk", "MAJOR",
     "Contrast-induced nephropathy + metformin accumulation", "Hold 48 h before contrast; restart after renal confirmation", "ACR 2024"],
    ["Metformin (MANKIND-M)", "Alcohol (chronic heavy use)", "Inhibits gluconeogenesis", "MODERATE",
     "Lactic acidosis risk in heavy drinkers", "Counsel on alcohol limits; monitor lactate if symptomatic", "WHO Guidelines"],
    ["Sitagliptin (DYNADUO)", "Digoxin", "P-gp inhibition by sitagliptin", "MINOR",
     "Slight increase in digoxin levels (~11%)", "Monitor digoxin; usually clinically insignificant", "NEJM 2023"],
    ["Semaglutide (GLP-1 RA)", "Warfarin / Anticoagulants", "Slowed gastric emptying", "MODERATE",
     "Altered anticoagulant absorption kinetics", "Frequent INR monitoring — first 4 weeks", "Lancet Diabetes 2024"],
    ["Semaglutide (GLP-1 RA)", "Oral Contraceptives", "Delayed gastric absorption", "MINOR",
     "Reduced Cmax of OC (AUC unchanged)", "Take OC same time daily; add barrier method first month", "Novo Nordisk PIB"],
    ["Canagliflozin", "Rifampicin", "CYP3A4 / UGT1A9 induction", "MAJOR",
     "60% reduction in canagliflozin AUC", "Increase dose to 300 mg; consider alternative SGLT2i", "J&J Prescribing Info"],
    ["Glimepiride (MANKIND-G)", "Fluconazole", "CYP2C9 inhibition", "MAJOR",
     "Up to 2-fold increase in glimepiride AUC — hypoglycaemia", "Reduce glimepiride dose 50%; increase glucose monitoring", "FDA Drug Label"],
    ["Glimepiride (MANKIND-G)", "Beta-blockers (non-selective)", "Masks hypoglycaemia symptoms", "MODERATE",
     "Tachycardia masked; sweating persists as warning", "Educate patient: sweating is primary hypoglycaemia signal", "ESC/EAS 2024"],
    ["Dapagliflozin", "Valproic acid", "UGT1A9 inhibition", "MINOR",
     "Slight increase in dapagliflozin exposure", "No dose adjustment; monitor for SGLT2i side effects", "EMA Review 2023"],
]

SEV_FILL = {"MAJOR": RED, "MODERATE": AMBER, "MINOR": GREEN}

for row_data in ROWS:
    ws.append(row_data)
    r = ws.max_row
    sev = row_data[3]
    for col in range(1, 8):
        c = ws.cell(row=r, column=col)
        c.alignment = WRAP
        if col == 4 and sev in SEV_FILL:
            c.fill = SEV_FILL[sev]
            c.font = Font(bold=True, size=10)

ws.column_dimensions["A"].width = 22
ws.column_dimensions["B"].width = 24
ws.column_dimensions["C"].width = 28
ws.column_dimensions["D"].width = 11
ws.column_dimensions["E"].width = 32
ws.column_dimensions["F"].width = 36
ws.column_dimensions["G"].width = 18

# Sheet 2 — prescribing trends
ws2 = wb.create_sheet("Prescribing Trends India 2024")
ws2.append(["Specialty", "% on SGLT2i", "% on GLP-1 RA", "% on Metformin", "Top Combination"])
for row in [
    ["Diabetology",      68, 42, 95, "Metformin + Empagliflozin"],
    ["Cardiology",       44, 28, 71, "Empagliflozin + ACEi/ARB"],
    ["Nephrology",       52, 18, 64, "Empagliflozin + ARB"],
    ["General Medicine", 28, 12, 88, "Metformin + Sitagliptin"],
    ["Endocrinology",    71, 55, 93, "Metformin + Semaglutide"],
    ["Gynaecology",       8,  3, 61, "Metformin (PCOS) + OCP"],
]:
    ws2.append(row)

wb.save(str(OUT / "Diabetes_Drug_Interaction_Database_2024.xlsx"))
print("[OK] Diabetes_Drug_Interaction_Database_2024.xlsx")


# =============================================================================
# FILE 3  GLP-1 Real-World Evidence  (.txt)
# =============================================================================
GLP1_TXT = """\
MANKIND PHARMA -- INTERNAL REPORT
GLP-1 Receptor Agonists: Real-World Evidence from Indian Diabetology Centres
Prepared by: Medical Affairs & HEOR Team  |  November 2024  |  STRICTLY CONFIDENTIAL
==================================================================================

EXECUTIVE SUMMARY
-----------------
24-month real-world data from 14 tertiary-care diabetology centres across India
(n=3,240 T2DM patients initiated on GLP-1 receptor agonists, Jan 2022 - Dec 2023).
Molecules covered: semaglutide (oral + injectable), dulaglutide, liraglutide, exenatide.

PATIENT DEMOGRAPHICS
--------------------
  Mean age                : 52.4 +/- 9.8 years
  Gender (M/F)            : 58% / 42%
  Mean BMI                : 29.6 +/- 5.2 kg/m2
  Mean baseline HbA1c     : 9.1 +/- 1.4%
  Mean T2DM duration      : 8.3 years
  Concomitant CVD         : 34%
  CKD (eGFR 30-60)        : 18%
  Top cities              : Mumbai (21%), Delhi (18%), Bangalore (14%), Chennai (11%)

EFFICACY AT 12 MONTHS
----------------------
  HbA1c reduction         : -1.6% semaglutide inj  |  -1.1% dulaglutide
  Body weight change      : -4.2 kg semaglutide inj  |  -3.1 kg dulaglutide
  Responders (HbA1c <7%)  : 48% semaglutide  |  38% dulaglutide  |  29% liraglutide
  Systolic BP reduction   : -5.1 mmHg (all GLP-1 RAs combined)
  LDL-C reduction         : -8.2 mg/dL (semaglutide inj)
  MACE (CVD sub-group)    : 22% reduction vs matched SoC cohort

SAFETY PROFILE (INDIA)
-----------------------
  GI side effects (nausea/vomiting)  : 28% -- mostly transient, weeks 1-8
  Discontinuation due to GI AE       : 6.8%
  Hypoglycaemia, mild                 : 4.2% monotherapy; 14.1% with sulphonylurea
  Injection-site reactions            : 2.1%
  Pancreatitis (confirmed)            : 0.04% (4 cases in 3,240 -- baseline rate)
  No DKA, no thyroid C-cell tumours (ongoing surveillance)

KOL INSIGHTS -- 48 INTERVIEWS
-------------------------------
Top reasons for initiating GLP-1 RA:
  1. High HbA1c despite dual OAD therapy            (82% KOLs)
  2. Weight management alongside glucose control     (74%)
  3. CVD/ASCVD risk reduction                        (68%)
  4. Patient request after media coverage of Ozempic (31%)

Top prescribing barriers:
  1. Cost (semaglutide inj: INR 3,500-5,200/month)  (91% KOLs)
  2. Patient preference for oral route               (44%)
  3. GI tolerability concerns                        (38%)
  4. Lack of awareness on renal-dosing flexibility   (22%)

MANKIND MARKET OPPORTUNITY
---------------------------
  Internal data: 12-14% market share opportunity in GLP-1 RA segment by 2026
  Key target segments:
    - T2DM + BMI >27 kg/m2 (weight-loss narrative)
    - T2DM + CVD (CV protection narrative)
    - High-HbA1c patients in Tier-B cities (lower competitive density)
  Flagship: MANKIND-SEMA (semaglutide 0.5 mg / 1 mg inj) -- Phase III expected Q2 2025

MSL RECOMMENDATIONS
--------------------
  1. Lead with weight + CV benefit -- not just HbA1c
  2. Use SUSTAIN-6 and SELECT trial data for cardiologist targeting
  3. Patient support programme to address cost barrier
  4. Emphasise oral semaglutide (Rybelsus equivalent) for injection-averse
  5. GLP-1 + SGLT2i combination pitch for CKD/HF patients

KEY REFERENCES
--------------
  SUSTAIN-6 Trial, NEJM 2016; Marso et al.
  SELECT Cardiovascular Outcomes Trial, NEJM 2023; Lincoff et al.
  STEP-5 Extension, Diabetes Care 2024; Garvey et al.
  India Real-World HEOR Study -- Mankind Pharma Internal Data 2024 (unpublished)
  ADA Standards of Care in Diabetes, Diabetes Care 2024

[END OF REPORT -- DO NOT DISTRIBUTE EXTERNALLY]
"""

(OUT / "GLP1_Real_World_India_Evidence_INTERNAL.txt").write_text(GLP1_TXT.strip(), encoding="utf-8")
print("[OK] GLP1_Real_World_India_Evidence_INTERNAL.txt")


# =============================================================================
# FILE 4  PCOS Management Protocol  (.docx)
# =============================================================================
doc2 = Document()
doc2.core_properties.author = "Dr. Sunanda Rao, Medical Affairs, Mankind Pharma"

t2 = doc2.add_heading("PCOS Management Protocol -- India 2024", level=1)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc2.add_paragraph("Version: 3.2  |  Reviewed: September 2024  |  Next Review: September 2025")
doc2.add_paragraph("Authors: Dr. Sunanda Rao (Endocrinology) | Dr. Kavita Pillai (Gynaecology) | Mankind Medical Affairs")
doc2.add_paragraph("Classification: INTERNAL CLINICAL PROTOCOL -- MR TRAINING USE ONLY")
doc2.add_paragraph("")

doc2.add_heading("1. Diagnostic Criteria (Rotterdam 2003 -- India Adapted)", level=2)
doc2.add_paragraph("Diagnosis: 2 of 3 criteria present, after excluding other causes:")
for c in [
    "Oligo/anovulation: <8 menstrual cycles/year OR cycle length >35 days",
    "Clinical hyperandrogenism (hirsutism, acne, alopecia) OR raised total/free testosterone",
    "PCOM on ultrasound: >=20 follicles per ovary (2.5-9 mm) OR ovarian volume >10 mL",
]:
    doc2.add_paragraph(c, style="List Bullet")
doc2.add_paragraph(
    "India-specific note: AMH >4.0 ng/mL supports diagnosis. Lean PCOS (BMI <23) accounts for "
    "20-30% of Indian cases -- do not dismiss based on BMI alone."
)

doc2.add_heading("2. First-Line Treatment Algorithm", level=2)
doc2.add_paragraph("Step 1: Lifestyle Modification (ALL patients -- mandatory)")
for item in [
    "Weight loss 5-10% body weight if BMI >23 kg/m2 (Indian cut-off)",
    "Mediterranean-style diet + 150 min/week aerobic activity",
    "Expected: menstrual regularity restoration in 30-40% of overweight patients",
]:
    doc2.add_paragraph(item, style="List Bullet")

doc2.add_paragraph("")
doc2.add_paragraph("Step 2A: Menstrual Regulation (fertility NOT desired)")
for item in [
    "Combined OCP (EE 20-30 mcg + desogestrel/drospirenone): first choice",
    "MANKIND-DRSP (drospirenone 3 mg + EE 20 mcg): preferred for anti-androgenic effect",
    "Cyclic progesterone (MPA 10 mg x 10 days/month): if OCP contraindicated",
]:
    doc2.add_paragraph(item, style="List Bullet")

doc2.add_paragraph("")
doc2.add_paragraph("Step 2B: Insulin Resistance Management")
for item in [
    "Metformin 500 mg BD (titrate to 1000 mg BD over 4 weeks): first choice in IR-PCOS",
    "MANKIND-M (Metformin SR): preferred for better GI tolerability",
    "Inositol (myo-inositol 4 g/day, 40:1 myo:d-chiro ratio): strong alternative in lean PCOS",
    "Expected: HbA1c reduction 0.5-0.8%; menstrual regularity improvement 40-50%",
]:
    doc2.add_paragraph(item, style="List Bullet")

doc2.add_heading("3. Fertility Management (ovulation induction)", level=2)
for item in [
    "Letrozole 2.5-5 mg D2-D6: FIRST LINE (supersedes clomiphene per ESHRE 2023)",
    "Clomiphene citrate 50-150 mg D2-D6: second line; max 6 cycles",
    "Add metformin if BMI >25 kg/m2 or HOMA-IR >2.5 (improves ovulation rate 15-20%)",
    "Refer for IVF/IUI after 6 failed ovulation induction cycles",
]:
    doc2.add_paragraph(item, style="List Bullet")

doc2.add_heading("4. Long-Term Metabolic Monitoring (Annual)", level=2)
for item in [
    "Fasting glucose + 2h OGTT (T2DM risk 3-7x higher in PCOS)",
    "Fasting lipid profile (dyslipidaemia in 70% of PCOS)",
    "Thyroid function TSH (hypothyroidism 2-3x more prevalent in PCOS)",
    "Blood pressure (HTN risk 3x higher by age 40-50)",
    "Endometrial thickness USS if >12 months amenorrhoea",
]:
    doc2.add_paragraph(item, style="List Bullet")

doc2.add_heading("5. Mankind Portfolio for PCOS", level=2)
doc2.add_paragraph(
    "MANKIND-DRSP (COCP)  |  MANKIND-M (Metformin SR)  |  MANKIND-LET (Letrozole). "
    "Combined addressable market: ~INR 340 crore in PCOS segment (FY25 estimate). "
    "MR pitch priority: gynaecologists and endocrinologists in metros + Tier-1 cities."
)

doc2.save(str(OUT / "PCOS_Management_Protocol_India_2024.docx"))
print("[OK] PCOS_Management_Protocol_India_2024.docx")


# =============================================================================
# Summary
# =============================================================================
print()
print("[OK] All 4 research documents created ->", str(OUT))
print()
for f in sorted(OUT.iterdir()):
    kb = f.stat().st_size // 1024
    print(f"     {f.suffix.upper()[1:]:5s}  {kb:>4d} KB   {f.name}")
print()
print("Agent Alpha will read these files automatically via the local_docs tool.")
print("Drop any additional .txt / .docx / .xlsx / .pdf files into the Research/ folder")
print("and Alpha will pick them up on the next pipeline run.")
